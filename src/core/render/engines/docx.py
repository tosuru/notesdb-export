"""
DOCX レンダリングエンジン (DocxRenderer v1.1)

- BaseRenderer を継承
- normalized.json の runs 配列を python-docx を使用して DOCX 文書に変換します。
- スタイル (太字、斜体、サイズ、色、配置、インデント、リスト) を適用します。
- 画像 (img) を埋め込みます。
- 添付ファイル (attachmentref) への相対パスリンクを生成します。
- テーブル (table) を生成し、セル結合や背景色も可能な限り反映します。
- v1.1: BaseRenderer の変更に対応 (doc_path 入力)、相対パスリンク実装
"""

from __future__ import annotations
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional
from io import BytesIO

try:
    from docx import Document
    from docx.document import Document as DocumentObject  # 型ヒント用
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run
    from docx.table import Table, _Cell
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import docx as python_docx  # For RELATIONSHIP_TYPE
    HAVE_DOCX_LIBS = True
except ImportError as e:
    HAVE_DOCX_LIBS = False
    _import_error_message = (
        "python-docx is required for DOCX rendering. "
        f"Install with: pip install python-docx. Original error: {e}"
    )

# --- common.py から BaseRenderer をインポート ---
try:
    # パッケージとして実行される場合
    from ..common import BaseRenderer, Style, _merge_styles  # _merge_stylesも利用
except ImportError:
    # スクリプトとして直接実行される場合など
    import sys
    common_dir = Path(__file__).resolve().parent.parent
    if str(common_dir) not in sys.path:
        sys.path.insert(0, str(common_dir))
    from common import BaseRenderer, Style, _merge_styles  # type: ignore

# --- Utilities (potentially move to utils module later) ---
from docx.oxml.exceptions import InvalidXmlError  # エラー処理用


def parse_color(color_str: Optional[str]) -> Optional[RGBColor]:
    """ #RGB 形式の文字列を RGBColor に変換 """
    if not color_str or not color_str.startswith('#') or len(color_str) != 7:
        return None
    try:
        r = int(color_str[1:3], 16)
        g = int(color_str[3:5], 16)
        b = int(color_str[5:7], 16)
        return RGBColor(r, g, b)
    except ValueError:
        return None


def parse_pt_size(size_str: Optional[str]) -> Optional[Pt]:
    """ '12pt' 形式の文字列を Pt に変換 """
    if not size_str or not size_str.lower().endswith('pt'):
        return None
    try:
        return Pt(float(size_str[:-2]))
    except ValueError:
        return None


def parse_length(length_str: Optional[str], unit='inches') -> Optional[Any]:
    """ '1.5in' 形式の文字列を指定単位 (Inches or Pt) に変換 """
    if not length_str:
        return None
    length_str = length_str.lower()
    val = None
    try:
        if length_str.endswith('in'):
            val = float(length_str[:-2])
            return Inches(val) if unit == 'inches' else Pt(val * 72)
        elif length_str.endswith('pt'):
            val = float(length_str[:-2])
            return Pt(val) if unit == 'pt' else Inches(val / 72)
        # 他の単位 (cm, mm) のサポートを追加する場合はここに記述
    except ValueError:
        pass
    logger.warning(f"Could not parse length '{length_str}' to {unit}")
    return None


def set_cell_background(cell: _Cell, color_str: Optional[str]):
    """ セルの背景色を設定 """
    color = parse_color(color_str)
    if not color:
        return
    try:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color.rgb_hex)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    except Exception as e:
        logger.warning(f"Failed to set cell background color {color_str}: {e}")


# --- Logger setup ---
logger = logging.getLogger(__name__)

# --- Main Renderer Class ---


class DocxRenderer(BaseRenderer):
    """
    normalized.json (runs) を python-docx を使用して DOCX 文書に変換するレンダラー。
    """

    def __init__(self, doc_path: Path):
        if not HAVE_DOCX_LIBS:
            raise ImportError(_import_error_message)

        # BaseRenderer の __init__ を呼び出す
        # これにより self.doc, self.attachment_dir, self.context, self.output が初期化される
        super().__init__(doc_path)

        # DocxRenderer 固有の状態
        self.document: DocumentObject = Document()
        self._current_paragraph: Optional[Paragraph] = None
        # デフォルトフォントの設定 (オプション - 日本語環境考慮)
        self._set_default_font()

    def _set_default_font(self, font_name="Meiryo UI"):  # 例: Meiryo UI
        """文書全体のデフォルトフォントを設定 (東アジア言語も考慮)"""
        try:
            style = self.document.styles['Normal']
            font = style.font
            font.name = font_name
            # 日本語などのフォント設定
            style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            logger.debug(f"Set default document font to: {font_name}")
        except Exception as e:
            logger.warning(f"Failed to set default font to {font_name}: {e}")

    # --- BaseRenderer 抽象メソッドの実装 ---

    def _init_output(self) -> BytesIO:
        """
        DOCX ファイルの内容を書き込むための BytesIO バッファを初期化します。
        """
        return BytesIO()

    def get_output(self) -> bytes:
        """
        構築された Document オブジェクトを BytesIO バッファに保存し、
        バイト列として返します。
        """
        if not isinstance(self.output, BytesIO):
            logger.error("Output buffer is not initialized correctly.")
            return b"Error: Output buffer not initialized."
        try:
            self.document.save(self.output)
            return self.output.getvalue()
        except Exception as e:
            logger.error(f"Failed to save DOCX document: {e}", exc_info=True)
            # エラー時でも部分的なバイト列を返す試み (ただし破損している可能性あり)
            # return self.output.getvalue()
            # より安全なエラーメッセージを返す
            return f"Error saving DOCX: {e}".encode('utf-8')

    def _render_header(self):
        """文書タイトル (Subject) とメタ情報 (UNID) を文書に追加します。"""
        subject = self.doc.get("fields", {}).get(
            "Subject", {}).get("value", self.doc_path.stem)
        meta = self.doc.get("meta", {})
        unid = meta.get("unid")

        self.document.add_heading(subject, level=1)
        if unid:
            p = self.document.add_paragraph()
            run = p.add_run(f"UNID: {unid}")
            run.italic = True
        self.document.add_paragraph()  # タイトル後のスペース

    def _render_footer(self):
        """フッター (ここでは何もしない)"""
        pass

    def _render_appendix(self):
        """付録データを文書の末尾にテーブルとして追加します。"""
        rows = self._build_appendix_rows_helper()
        if not rows:
            logger.debug("No appendix data to render.")
            return

        self.document.add_page_break()  # 付録は新しいページから
        self.document.add_heading("付録：その他のフィールド", level=1)
        if not rows:
            self.document.add_paragraph("（該当フィールドなし）")
            return

        try:
            # テーブル作成 (ヘッダー行 + データ行)
            table = self.document.add_table(rows=1, cols=3, style='Table Grid')
            table.autofit = True  # 自動調整を試みる

            # ヘッダー設定
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "フィールド名"
            hdr_cells[1].text = "型"
            hdr_cells[2].text = "プレビュー"
            # ヘッダー行を太字に (オプション)
            for cell in hdr_cells:
                if cell.paragraphs:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

            # データ行追加
            for r in rows:
                row_cells = table.add_row().cells
                row_cells[0].text = str(r.get("name", ""))
                row_cells[1].text = str(r.get("type", ""))
                # プレビュー内の改行を反映させる (add_paragraphを使う)
                preview_text = str(r.get("preview", ""))
                # 既存の段落をクリア (通常は1つだけ存在)
                # paragraphs[0].clear() のような直接的なクリアがないため
                row_cells[2].text = ""
                for i, line in enumerate(preview_text.splitlines()):
                    if i == 0:
                        row_cells[2].paragraphs[0].add_run(line)
                    else:
                        row_cells[2].add_paragraph(line)
                # 垂直方向の配置を上揃えに
                row_cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            logger.debug(f"Appendix table rendered with {len(rows)} rows.")

        except Exception as e:
            logger.error(
                f"Failed to render appendix table: {e}", exc_info=True)
            self.document.add_paragraph(f"[付録テーブルの描画中にエラー発生: {e}]")

    def _start_paragraph(self, par_style: Dict[str, Any]):
        """
        新しい段落 (Paragraph オブジェクト) を作成し、スタイルを適用します。
        現在の段落として self._current_paragraph に保持します。
        """
        # _ensure_paragraph_started から空の par_style で呼ばれることも考慮
        par_style = par_style or {}

        # 新しい段落を追加し、現在の段落として設定
        self._current_paragraph = self.document.add_paragraph()
        self.context.paragraph_started = True  # BaseRenderer の状態更新

        # スタイル適用
        p_format = self._current_paragraph.paragraph_format

        # 1. 配置 (Alignment)
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "full": WD_ALIGN_PARAGRAPH.JUSTIFY,  # DXL 'full' -> JUSTIFY
        }
        alignment = align_map.get(str(par_style.get('align', '')).lower())
        if alignment is not None:
            p_format.alignment = alignment

        # 2. インデント (Left Margin)
        left_margin = parse_length(par_style.get('leftmargin'), 'inches')
        if left_margin is not None:
            p_format.left_indent = left_margin

        # 3. リストスタイル (簡易版 - Word の組み込みスタイル名を使用)
        list_state = self.context.list_state  # common.pyで更新された状態
        if list_state['level'] > 0:
            list_type = list_state.get('type')
            style_name = 'List Paragraph'  # デフォルト
            # TODO: list_typeに応じて 'List Bullet', 'List Number' などを設定
            # NotesのリストタイプとWordのスタイル名のマッピングが必要
            if list_type in ('bullet', 'uncheck', 'square'):
                style_name = 'List Bullet'
            elif list_type in ('number', 'alphaupper', 'alphalower', 'romanupper', 'romanlower'):
                style_name = 'List Number'

            try:
                self._current_paragraph.style = self.document.styles[style_name]
                # TODO: ネストレベルに応じたインデント調整が必要な場合がある
                # p_format.left_indent += Inches(0.25 * (list_state['level'] - 1)) # 例
            except KeyError:
                logger.warning(
                    f"List style '{style_name}' not found. Using default paragraph.")
                # フォールバック
                self._current_paragraph.style = self.document.styles['List Paragraph']

        # 4. 段落後スペース (Space After)
        space_after = parse_length(
            par_style.get('spaceafter'), 'pt')  # ポイント単位で
        if space_after is not None:
            p_format.space_after = space_after
        # else:
            # デフォルトの段落後スペースを設定 (必要に応じて)
            # p_format.space_after = Pt(6)

        # 5. Word組み込みの段落スタイル (parstyle) - 実験的
        # parstyle_name = par_style.get('parstyle')
        # if parstyle_name and parstyle_name in self.document.styles:
        #     try:
        #         self._current_paragraph.style = self.document.styles[parstyle_name]
        #     except KeyError:
        #         logger.warning(f"Paragraph style '{parstyle_name}' not found.")

        logger.debug(
            f"Started paragraph with style: {par_style}, List: {list_state}")

    def _finalize_paragraph(self):
        """
        現在の段落への参照をクリアします。
        (python-docxでは段落オブジェクト自体はDocumentに残る)
        """
        # テーブルなどのブロック要素処理の前に現在の段落を確定させる
        # 特に処理は不要だが、状態をリセット
        self._current_paragraph = None
        # self.context.paragraph_started は BaseRenderer 側で管理される
        # logger.debug("Finalized paragraph.")

    def _ensure_paragraph_started(self):
        """
        まだ段落 (_current_paragraph) が開始されていない場合、
        デフォルトのスタイルで新しい段落を開始します。
        """
        if self._current_paragraph is None:
            logger.debug("No current paragraph, ensuring default paragraph.")
            # コンテキストから現在のスタイルを取得して開始
            self._start_paragraph(self.context.current_par_style)

    def _append_run(self, text: str) -> Optional[Run]:
        """現在の段落に Run を追加して返す。段落がなければエラーログ。"""
        self._ensure_paragraph_started()
        if self._current_paragraph:
            try:
                return self._current_paragraph.add_run(text)
            except Exception as e:
                logger.error(
                    f"Error adding run with text '{text[:50]}...': {e}", exc_info=True)
                return None
        else:
            logger.error("Cannot add run: _current_paragraph is None.")
            return None

    # --- ハンドラメソッド ---

    def _handle_text(self, run_data: Dict[str, Any]):
        """'text' トークンを処理し、スタイル付きの Run を現在の段落に追加します。"""
        text = run_data.get('text', '')
        if not text:
            return  # 空テキストは無視

        docx_run = self._append_run(text)
        if not docx_run:
            return  # Runの追加に失敗したら何もしない

        # スタイル適用
        s = run_data.get('s', [])
        a = run_data.get('a', {})
        font = docx_run.font

        if 'b' in s:
            docx_run.bold = True
        if 'i' in s:
            docx_run.italic = True
        if 'u' in s:
            docx_run.underline = True
        if 's' in s or 'strike' in s:
            font.strike = True
        # if 'mono' in s: font.name = 'Courier New' # 等幅フォント指定 (要フォント存在確認)

        color = parse_color(a.get('color'))
        if color:
            font.color.rgb = color
        bgcolor = parse_color(a.get('bgcolor'))  # 背景色はRunレベルでは難しい
        if bgcolor:
            # ハイライト色として設定 (近似)
            try:
                # 色をWD_COLOR_INDEXにマッピングするか、カスタムカラーを設定
                # ここでは簡易的に黄色に固定
                if bgcolor.rgb_hex.lower() == 'ffff00':  # 黄色の場合
                    font.highlight_color = python_docx.enum.text.WD_COLOR_INDEX.YELLOW
                else:  # 他の色は無視またはログ出力
                    logger.debug(
                        f"Ignoring background color {a.get('bgcolor')} for run.")
            except AttributeError:  # highlight_color がないバージョンの場合
                logger.warning(
                    "Highlight color not supported in this python-docx version.")

        size = parse_pt_size(a.get('size'))
        if size:
            font.size = size

        script = a.get('script') or next(
            (fx for fx in a.get('fx', []) if fx in ['super', 'sub']), None)
        if script == 'super':
            font.superscript = True
        if script == 'sub':
            font.subscript = True

        fx = a.get('fx', [])
        if 'shadow' in fx:
            font.shadow = True
        if 'emboss' in fx:
            font.emboss = True
        if 'imprint' in fx:
            font.imprint = True  # extrude に近い効果
        # if 'outline' in fx: font.outline = True # バージョンによる

    def _handle_link(self, run_data: Dict[str, Any]):
        """'link' トークンを処理し、ハイパーリンクを現在の段落に追加します。"""
        label = run_data.get('label', run_data.get('href', ''))
        href = run_data.get('href', '#')
        notes_meta = run_data.get('notes')  # Notesリンク情報

        self._ensure_paragraph_started()
        if not self._current_paragraph:
            logger.error("Cannot add link: _current_paragraph is None.")
            return

        # Notesリンクの場合、情報を追記 (オプション)
        if notes_meta:
            display_text = f"{label} (Notes Link)"  # または notes_meta の内容を使う
        else:
            display_text = label

        try:
            # リンク色と下線を指定してハイパーリンクを追加
            self._add_hyperlink(self._current_paragraph, href,
                                display_text, color="0000FF", underline=True)
        except Exception as e:
            logger.error(
                f"Failed to add hyperlink '{href}': {e}", exc_info=True)
            # エラー時は通常のテキストとして追加 (フォールバック)
            run = self._append_run(f"{display_text} [{href}]")
            if run:
                run.italic = True

    def _handle_img(self, run_data: Dict[str, Any]):
        """'img' トークンを処理し、画像を現在の段落に埋め込みます。"""
        alt_text = run_data.get('alt', 'image')
        # content_path (相対パス) を使用
        content_path_rel = run_data.get('content_path')

        if not content_path_rel:
            logger.warning(f"Image run has no 'content_path'. Alt: {alt_text}")
            run = self._append_run(f"[画像(パスなし): {alt_text}]")
            if run:
                run.italic = True
            return

        # 絶対パスに解決
        img_path = self.resolve_attachment_path(content_path_rel)

        if img_path and img_path.is_file():
            self._ensure_paragraph_started()
            if not self._current_paragraph:
                logger.error("Cannot add image: _current_paragraph is None.")
                return
            try:
                # 画像を追加 (幅を指定可能, 例: 6インチ)
                # self._current_paragraph.add_run().add_picture(str(img_path), width=Inches(6.0))
                # サイズ指定なしで追加
                self._current_paragraph.add_run().add_picture(str(img_path))
                logger.debug(f"Embedded image: {img_path}")
            except (FileNotFoundError, InvalidXmlError) as img_err:  # 画像読み込みエラー処理
                logger.error(f"Failed to embed image '{img_path}': {img_err}")
                run = self._append_run(f"[画像(挿入失敗): {alt_text}]")
                if run:
                    run.italic = True
            except Exception as img_err:  # その他の予期せぬエラー
                logger.error(
                    f"Unexpected error embedding image '{img_path}': {img_err}", exc_info=True)
                run = self._append_run(f"[画像(不明なエラー): {alt_text}]")
                if run:
                    run.italic = True
        else:
            logger.warning(
                f"Image file not found or path is invalid: {img_path} (from {content_path_rel})")
            run = self._append_run(f"[画像(ファイル無し): {alt_text}]")
            if run:
                run.italic = True

    def _handle_table(self, run_data: Dict[str, Any]):
        """'table' トークンを処理し、文書にテーブルを追加します。セル結合も扱います。"""
        # テーブルはブロック要素なので、現在の段落を終了
        self._finalize_paragraph()

        rows_data = run_data.get('rows', [])
        if not rows_data:
            logger.warning("Table run found with no rows.")
            return

        # 最初の行から列数を推定 (可変長の場合があるため注意)
        num_cols = 0
        if rows_data[0].get('cells'):
            num_cols = len(rows_data[0]['cells'])
        if num_cols == 0:
            logger.warning(
                "Table row data exists, but the first row has no cells.")
            return

        logger.debug(
            f"Starting DOCX table rendering: {len(rows_data)} rows, estimated {num_cols} cols.")

        try:
            # テーブル作成 (行数=rows_dataの長さ, 列数=最初の行のセル数)
            # テーブルスタイルの適用 (オプション)
            table = self.document.add_table(
                rows=len(rows_data), cols=num_cols, style='Table Grid')
            table.autofit = True  # 自動調整

            # セル結合情報を保持する (結合開始セル座標 -> True)
            merged_cells = {}  # Dict[Tuple[int, int], bool]

            for r_idx, row in enumerate(rows_data):
                cells_data = row.get('cells', [])
                table_row = table.rows[r_idx]
                actual_c_idx = 0  # 実際のテーブル列インデックス (結合セルをスキップするため)

                for c_idx_data, cell_data in enumerate(cells_data):
                    # --- 結合セルをスキップ ---
                    while (r_idx, actual_c_idx) in merged_cells:
                        actual_c_idx += 1
                        if actual_c_idx >= num_cols:  # 行の範囲外になったら抜ける
                            break
                    if actual_c_idx >= num_cols:
                        logger.warning(
                            f"Skipping cell data index {c_idx_data} as it exceeds table column count due to merged cells.")
                        continue

                    # --- 現在のセルを取得 ---
                    try:
                        current_cell = table_row.cells[actual_c_idx]
                    except IndexError:
                        logger.error(
                            f"Cell index out of range: row={r_idx}, col={actual_c_idx} (num_cols={num_cols}). Skipping cell.")
                        actual_c_idx += 1  # とりあえずインデックスを進める
                        continue

                    # --- セル結合処理 ---
                    colspan = cell_data.get('colspan', 1)
                    rowspan = cell_data.get('rowspan', 1)
                    is_merged = False
                    if colspan > 1 or rowspan > 1:
                        try:
                            end_cell = table.cell(
                                r_idx + rowspan - 1, actual_c_idx + colspan - 1)
                            merged_cell_obj = current_cell.merge(end_cell)
                            is_merged = True
                            # 結合されたエリアのセルを記録して後でスキップ
                            for mr in range(r_idx, r_idx + rowspan):
                                for mc in range(actual_c_idx, actual_c_idx + colspan):
                                    if (mr, mc) != (r_idx, actual_c_idx):
                                        merged_cells[(mr, mc)] = True
                            # 結合後の処理対象セルは左上のセル (merged_cell_obj)
                            current_cell = merged_cell_obj
                        except IndexError:
                            logger.warning(
                                f"Cell merge failed for cell ({r_idx},{actual_c_idx}): colspan={colspan}, rowspan={rowspan}. Out of table bounds?")
                        except Exception as merge_err:  # 他の予期せぬエラー
                            logger.error(
                                f"Error merging cell ({r_idx},{actual_c_idx}): {merge_err}", exc_info=False)

                    # --- セル内容のレンダリング ---
                    # セルの最初の段落を取得 (結合した場合も同様)
                    # 既存の内容をクリアする必要があるか確認 (通常は不要)
                    if current_cell.paragraphs:
                        cell_paragraph = current_cell.paragraphs[0]
                        # 既存の Run をクリア (必要であれば)
                        # for run in cell_paragraph.runs:
                        #     p = cell_paragraph._element
                        #     p.remove(run._element)
                        cell_paragraph.text = ""  # 簡易的なクリア
                    else:
                        # 段落がない場合は追加 (通常は発生しないはず)
                        cell_paragraph = current_cell.add_paragraph()

                    # ヘルパー関数でセル内容をレンダリング
                    self._render_cell_content(
                        cell_paragraph, cell_data.get('runs', []))

                    # --- セルスタイル適用 (背景色) ---
                    cell_style = cell_data.get('style', {})
                    if isinstance(cell_style, dict):
                        inner_style = cell_style.get(
                            'style', {})  # 'style' キーがネストしている場合
                        bgcolor = inner_style.get('bgcolor')
                        if bgcolor:
                            set_cell_background(current_cell, bgcolor)

                    # --- 次の実際の列インデックスへ ---
                    actual_c_idx += colspan  # 結合した分だけ進める

            # タブ付き表のラベル (Markdown同様、表現が難しい)
            if run_data.get('attributes', {}).get('rowdisplay') == 'tabs':
                logger.warning(
                    "Tabbed table cannot be fully represented in DOCX. Rendering as a normal table.")
                # TODO: ラベルを行の先頭に追加するなどの処理？

            # テーブルの後に新しい段落を強制的に開始するための準備
            # 次のrun処理で _ensure_paragraph_started が呼ばれるように
            self.context.paragraph_started = False
            logger.debug(f"Finished DOCX table rendering.")

        except Exception as e:
            logger.error(f"Error rendering DOCX table: {e}", exc_info=True)
            self.document.add_paragraph(f"[テーブル描画エラー: {e}]")
            self.context.paragraph_started = False  # エラー後も新しい段落を開始

    def _render_cell_content(self, paragraph: Paragraph, runs: List[Dict[str, Any]]):
        """
        指定された Paragraph オブジェクトに、セル内の runs をレンダリングします。
        一時的に self._current_paragraph を設定して _process_runs を呼び出します。
        """
        original_paragraph = self._current_paragraph
        self._current_paragraph = paragraph  # 現在の段落を一時的にセル内の段落に設定
        try:
            # runs を処理 (これが text, link, img などを処理)
            self._process_runs(runs)
        finally:
            self._current_paragraph = original_paragraph  # 元の段落に戻す

    def _handle_attachmentref(self, run_data: Dict[str, Any]):
        """
        'attachmentref' トークンを処理し、添付ファイルへの相対パスリンクを
        現在の段落に追加します。
        """
        name = run_data.get('name', 'file')
        display_name = run_data.get('displayname', name)
        # content_path (相対パス) を使用
        content_path_rel = run_data.get('content_path')

        self._ensure_paragraph_started()
        if not self._current_paragraph:
            logger.error(
                "Cannot add attachment link: _current_paragraph is None.")
            return

        if content_path_rel:
            try:
                # add_hyperlink 関数 (移植したもの) を使用
                # URL には相対パスをそのまま渡す
                # Word が開く際に基準ディレクトリからの相対パスとして解釈することを期待
                # スラッシュは Word のハイパーリンク形式に合わせておく (バックスラッシュでも動く場合あり)
                relative_url = content_path_rel.replace('\\', '/')
                self._add_hyperlink(self._current_paragraph, relative_url,
                                    display_name, color="0000FF", underline=True)
                logger.debug(
                    f"Added relative hyperlink for attachment: {display_name} -> {relative_url}")
            except Exception as e:
                logger.error(
                    f"Failed to add relative hyperlink for '{display_name}' -> '{content_path_rel}': {e}", exc_info=True)
                # エラー時は通常のテキストとして追加 (フォールバック)
                run = self._append_run(f"[添付ファイル(リンク失敗): {display_name}]")
                if run:
                    run.italic = True
        else:
            logger.warning(
                f"AttachmentRef run has no 'content_path'. Name: {name}")
            # content_path がなければテキストのみ
            run = self._append_run(f"[添付ファイル: {display_name}]")
            if run:
                run.italic = True

    def _handle_br(self, run_data: Dict[str, Any]):
        """
        'br' トークンを処理し、新しい段落を開始します。
        現在の段落は終了させます。
        """
        self._finalize_paragraph()
        # 新しい段落を開始 (スタイルは現在のコンテキストを引き継ぐ)
        self._start_paragraph(self.context.current_par_style)
        logger.debug("Processed 'br' by starting a new paragraph.")

    def _handle_unknown(self, run_data: Dict[str, Any]):
        """不明なトークンタイプを処理します (フォールバック)。"""
        run_type = run_data.get('t', 'unknown')
        logger.warning(
            f"Encountered unknown run type: '{run_type}'. Rendering as placeholder text.")
        run = self._append_run(f"[Unknown Run: {run_type}]")
        if run:
            run.italic = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 赤色

    # --- Relative Hyperlink Helper (Ported) ---
    def _add_hyperlink(self, paragraph: Paragraph, url: str, text: str, color: str = "0000FF", underline: bool = True):
        """
        指定された段落にハイパーリンクを追加します (相対パス対応)。
        docx_builder 相対パス例.py から移植・修正。
        """
        part = paragraph.part
        # is_external=True のままで相対パスを渡す
        # Word が RELATIONSHIP_TYPE.HYPERLINK と TargetMode=External で
        # 相対パスを正しく解釈することを期待
        try:
            r_id = part.relate_to(
                url,
                python_docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                is_external=True
            )
        except AttributeError:
            logger.error(
                "Could not find RELATIONSHIP_TYPE. Maybe python-docx version issue?")
            # フォールバックとして通常のテキスト追加
            run = self._append_run(f"{text} [{url}]")
            if run:
                run.italic = True
            return None

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # ハイパーリンクのスタイル付き Run を作成
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # スタイル: Hyperlink (Word組み込みスタイル)
        # これにより、テーマに基づいた色や下線が適用されることが多い
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)

        # 明示的な色指定 (オプション)
        # if color:
        #     c = OxmlElement('w:color')
        #     c.set(qn('w:val'), color)
        #     rPr.append(c)

        # 明示的な下線指定 (オプション - 通常Hyperlinkスタイルに含まれる)
        # if underline:
        #     u = OxmlElement('w:u')
        #     u.set(qn('w:val'), 'single')
        #     rPr.append(u)

        new_run.append(rPr)

        # テキストノードを追加
        t_element = OxmlElement('w:t')
        t_element.text = text
        new_run.append(t_element)

        hyperlink.append(new_run)

        # 段落の XML 要素 (_p) にハイパーリンク要素を追加
        paragraph._p.append(hyperlink)

        return hyperlink

    # --- Appendix Helper ---

    def _build_appendix_rows_helper(self) -> List[Dict[str, Any]]:
        """ 付録テーブル用の行データを構築 (md.py とほぼ同じ) """
        # (md.py からコピー＆ペースト - 必要なら調整)
        rows = []
        fields = self.doc.get("fields", {}) or {}
        layout = self.doc.get("layout", {}) or {}
        allow = set(layout.get("primary_fields_allowlist", []) or [])
        used = set(layout.get("used_in_body", []) or [])
        logger.debug(f"Appendix build: allowlist={allow}, used_in_body={used}")
        processed_names = set()

        for name, meta in fields.items():
            if name.startswith('$') or name in allow or name in used or name in processed_names:
                continue

            ftype = meta.get("type", "unknown")
            val = meta.get("value")
            preview = ""
            try:
                if ftype == 'richtext':
                    preview = meta.get(
                        'text', '[RichText Body]')  # プレーンテキストを使用
                elif isinstance(val, list):
                    if val and all(isinstance(x, (str, int, float, bool, type(None))) for x in val):
                        preview = ", ".join(map(str, val[:5]))
                        if len(val) > 5:
                            preview += " ..."
                    else:
                        preview = f"[List of {len(val)} complex items]"
                elif isinstance(val, dict):
                    preview = "{...}"
                elif val is None:
                    preview = "[None]"
                else:
                    preview = str(val)

                max_len = 100  # DOCXではもう少し長くても良いかも
                if len(preview) > max_len:
                    preview = preview[:max_len] + " ..."

            except Exception as e:
                preview = f"[Error previewing: {e}]"
                logger.warning(
                    f"Error generating preview for field '{name}': {e}", exc_info=False)

            rows.append({"name": name, "type": ftype, "preview": preview})
            processed_names.add(name)

        rows.sort(key=lambda r: r["name"].lower())
        logger.debug(f"Generated {len(rows)} appendix rows.")
        return rows

    def get_output_on_error(self, e: Exception) -> bytes:
        """レンダリングエラー発生時にエラー情報を埋め込んだ DOCX を返す"""
        logger.error(f"Render error occurred: {e}", exc_info=True)
        try:
            # 既存の文書にエラー情報を追記
            self.document.add_page_break()
            self.document.add_heading("Render Error", level=1)
            p = self.document.add_paragraph()
            run = p.add_run(f"{type(e).__name__}: {e}")
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 赤色
            # スタックトレースも追加 (オプション)
            import traceback
            tb_lines = traceback.format_exception(
                type(e), e, e.__traceback__, limit=5)
            p_tb = self.document.add_paragraph()
            run_tb = p_tb.add_run("Traceback (partial):\n" + "".join(tb_lines))
            run_tb.font.size = Pt(8)

            # エラー情報を含む文書を保存
            error_output = BytesIO()
            self.document.save(error_output)
            return error_output.getvalue()

        except Exception as save_err:
            # エラー追記・保存自体に失敗した場合
            logger.error(
                f"Failed to save DOCX even with error info: {save_err}")
            return f"Render Error: {e}\nFailed to save document with error details.".encode('utf-8')


# --- 簡易テストコード ---
if __name__ == "__main__":
    logging.basicConfig(
        level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(name)s - %(message)s')

    script_dir = Path(__file__).resolve().parent
    try:
        project_root = script_dir.parent.parent
        test_data_dir = project_root / "test_data"
        if not test_data_dir.is_dir():
            test_data_dir = script_dir
            logger.warning(
                f"Directory '{project_root / 'test_data'}' not found. Using '{script_dir}'.")
        # ここでテストに使うJSONファイルを指定
        test_json_path = test_data_dir / \
            "670309E2B88C5E9649258D2C000ADBDE.normalized.1.4.4-dev_initial.json"

    except Exception as e:
        print(f"Error constructing test JSON path: {e}")
        sys.exit(1)

    if test_json_path.exists():
        print(f"--- Running DOCX Render Test for: {test_json_path} ---")
        try:
            # attachments ディレクトリが JSON と同じ場所にあると仮定
            # テスト用に content_path をダミーで設定 (通常は attachment extractor が行う)
            import json
            with open(test_json_path, 'r', encoding='utf-8') as f:
                doc_data = json.load(f)

            # ダミーの content_path を設定 (attachments ディレクトリを想定)
            if 'attachments' in doc_data and isinstance(doc_data['attachments'], list):
                for i, att in enumerate(doc_data['attachments']):
                    if 'name' in att and att.get('content_path') is None:
                        # 拡張子がない場合があるかもしれないので注意
                        fname = Path(att['name'])
                        # 例: attachments/image_0.png
                        att['content_path'] = f"attachments/{fname.stem}_{i}{fname.suffix}"
                        logger.info(
                            f"Assigning dummy content_path: {att['content_path']} for {att['name']}")

            # 変更を一時ファイルに保存してテスト (元のファイルを変更しない)
            temp_json_path = test_json_path.with_name(
                test_json_path.stem + "_temp_test.json")
            with open(temp_json_path, 'w', encoding='utf-8') as f:
                json.dump(doc_data, f, ensure_ascii=False, indent=2)

            # レンダラーを実行
            renderer = DocxRenderer(temp_json_path)
            output_bytes = renderer.render()

            # 出力先DOCXファイルのパス
            output_docx_path = test_json_path.parent / \
                (test_json_path.stem + "_render_test.docx")

            with open(output_docx_path, "wb") as f:
                f.write(output_bytes)
            print(
                f"+++ DOCX render test successful. Output saved to: {output_docx_path.resolve()} +++")

            # 一時ファイルを削除
            temp_json_path.unlink()

        except ImportError as import_err:
            print(
                f"!!! DOCX Render Test SKIPPED due to missing library: {import_err} !!!")
        except FileNotFoundError as fnf_err:
            print(
                f"!!! DOCX Render Test FAILED: File not found - {fnf_err} !!!")
            logger.error(
                f"Make sure the input JSON and potentially referenced attachment files exist.")
        except Exception as main_e:
            print(f"--- DOCX Render Test FAILED ---")
            logger.exception("Error during DOCX render test")
            # エラー時の出力もファイルに保存してみる (デバッグ用)
            try:
                if 'output_bytes' in locals() and isinstance(output_bytes, bytes):
                    error_docx_path = test_json_path.parent / \
                        (test_json_path.stem + "_render_error.docx")
                    with open(error_docx_path, "wb") as f_err:
                        f_err.write(output_bytes)
                    print(
                        f"--- Error output saved to: {error_docx_path.resolve()} ---")
            except Exception as save_err:
                print(f"Could not save error output file: {save_err}")

    else:
        print(f"!!! Test JSON file not found: {test_json_path.resolve()} !!!")
        print("Please create a test JSON file or adjust the path.")
